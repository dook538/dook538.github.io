from docx import Document

# 1. สร้างเอกสารใหม่
doc = Document()

# 2. เพิ่มหัวข้อ (Heading)
doc.add_heading('บันทึกการเรียนรู้ของครูดุ๊ก', 0)

# 3. เพิ่มย่อหน้า (Paragraph)
p = doc.add_paragraph('วันนี้ผมได้เริ่มต้นใช้งาน ')
p.add_run('Gemini CLI (เจม)').bold = True
p.add_run(' เพื่อช่วยงานต่างๆ ในเครื่อง Mac')

# 4. เพิ่มรายการ (List)
doc.add_heading('สิ่งที่ได้เรียนรู้วันนี้:', level=1)
doc.add_paragraph('การสร้างเว็บเพจเบื้องต้นด้วย HTML/CSS', style='List Bullet')
doc.add_paragraph('การใช้ Python สร้างรูปภาพ (.png)', style='List Bullet')
doc.add_paragraph('การจัดระเบียบหน้าจอ Desktop', style='List Bullet')
doc.add_paragraph('การใช้โค้ดสร้างไฟล์ Word (.docx)', style='List Bullet')

# 5. บันทึกไฟล์
output_path = '/Users/duke/Desktop/เอกสาร (Documents)/บันทึกการเรียนรู้_ครูดุ๊ก.docx'
doc.save(output_path)

print(f"สร้างไฟล์ Word สำเร็จแล้วที่: {output_path}")
